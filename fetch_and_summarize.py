import os
import asyncio
import logging
import json
from datetime import datetime, timedelta, timezone
from pathlib import Path
import requests
from telethon import TelegramClient, events
from docx import Document

# ================== LOGGING ==================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler()]
)
log = logging.getLogger(__name__)

# ================== CONFIGURATION (from environment variables) ==================
API_ID = int(os.environ["API_ID"])
API_HASH = os.environ["API_HASH"]
BOT_TOKEN = os.environ["BOT_TOKEN"]
GEMINI_KEY = os.environ["GEMINI_KEY"]
OWNER_CHAT_ID = int(os.environ["OWNER_CHAT_ID"])

# The list of channels for gathering news and the time for the final submission are specified right here
CHANNELS = ["news chenel wich you want", "news chanel wich you want"]
SEND_HOUR = int(os.environ.get("SEND_HOUR", 21))
SEND_MINUTE = int(os.environ.get("SEND_MINUTE", 0))
COLLECT_MINUTES_BEFORE = int(os.environ.get("COLLECT_MINUTES_BEFORE", 10))

BASE_DIR = Path(__file__).parent
SUBSCRIBERS_FILE = BASE_DIR / "subscribers.json"
SESSION_FILE = str(BASE_DIR / "news_parser_session")

client = TelegramClient(SESSION_FILE, API_ID, API_HASH)          # user session — reads channels
bot_client = TelegramClient("bot_sender_session", API_ID, API_HASH)  # bot session — sends out a summary


# ================== SUBSCRIBERS ==================
def load_subscribers() -> list[int]:
    if SUBSCRIBERS_FILE.exists():
        try:
            with open(SUBSCRIBERS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            log.error(f"Ошибка чтения файла подписчиков: {e}")

    subs = [OWNER_CHAT_ID]
    save_subscribers(subs)
    return subs


def save_subscribers(subs: list[int]) -> None:
    try:
        with open(SUBSCRIBERS_FILE, "w", encoding="utf-8") as f:
            json.dump(subs, f, ensure_ascii=False)
    except Exception as e:
        log.error(f"Ошибка записи файла подписчиков: {e}")


# ================== NEWS COLLECTION AND SUMMARY ==================
async def fetch_news_from_channels() -> str:
    """Собирает текст всех постов за последние 24 часа из списка CHANNELS."""
    log.info("Начинаю сбор новостей из каналов за последние 24 часа...")
    collected_text = ""
    time_threshold = datetime.now(timezone.utc) - timedelta(days=1) 
    posts_count = 0

    for channel in CHANNELS:
        try:
            log.info(f"Сканирую канал: {channel}")
            async for message in client.iter_messages(channel):
                if not message.date or message.date < time_threshold:
                    break
                if message.text and len(message.text.strip()) > 10:
                    collected_text += f"\n--- Источник: {channel} ---\n{message.text}\n"
                    posts_count += 1
        except Exception as e:
            log.error(f"Не удалось собрать новости из {channel}: {e}")

    log.info(f"Собрано постов для анализа: {posts_count}")
    return collected_text


async def get_gemini_summary(text: str) -> str:
    """Суммаризирует собранный текст через Gemini API."""
    if not text.strip():
        return "За последние 24 часа не нашлось текстовых новостей в указанных каналах."

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={GEMINI_KEY}"
    prompt = (
        "Перед тобой сырой текст всех новостей за сегодня из нескольких крупных каналов. "
        "Твоя задача — выступить в роли шеф-редактора. Внимательно проанализируй весь массив данных, "
        "убери дубликаты, отсей второстепенный мусор и кликбейт. "
        "Сформируй из этого красивую, ёмкую, структурированную вечернюю сводку главных событий дня. "
        "Группируй логически, пиши интересно, используй смайлики для оформления и выделяй важное жирным шрифтом.\n\n"
        f"ТЕКСТ НОВОСТЕЙ:\n{text}"
    )
    payload = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(url, json=payload, timeout=600)
        response.raise_for_status()
        data = response.json()
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        log.error(f"Ошибка при обращении к Gemini API: {e}")
        return "Извините, не удалось сгенерировать сводку новостей через ИИ."


async def send_summary_to_subscribers(summary_text: str) -> None:
    """Формирует .docx со сводкой и рассылает его всем подписчикам от лица бота."""
    subs = load_subscribers()
    if not subs:
        log.info("Нет подписчиков для рассылки.")
        return

    date_str = datetime.now().strftime("%Y-%m-%d")
    filename = f"Сводка_новостей_{date_str}.docx"

    try:
        doc = Document()
        doc.add_heading(f"Вечерняя сводка новостей от {date_str}", level=1)
        for paragraph in summary_text.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        doc.save(filename)
        log.info(f"Документ {filename} создан.")
    except Exception as e:
        log.error(f"Ошибка при создании Word-файла: {e}")
        return

    log.info(f"Рассылаю файл {len(subs)} подписчикам...")

    for user_id in subs:
        try:
            await bot_client.send_file(
                user_id, filename,
                caption=f"📝 Ваш вечерний дайджест новостей готов ({date_str})!"
            )
        except Exception as e:
            log.error(f"Не удалось отправить файл пользователю {user_id}: {e}")


    if os.path.exists(filename):
        os.remove(filename)


# ================== COMMANDS AND PLANNER ==================
@bot_client.on(events.NewMessage(pattern="/start"))
async def start_command_handler(event):
    chat_id = event.chat_id
    subs = load_subscribers()
    if chat_id not in subs:
        subs.append(chat_id)
        save_subscribers(subs)
        log.info(f"Добавлен новый подписчик: {chat_id}")
    await event.respond("👋 Привет! Ты успешно подписался на вечернюю сводку новостей. Она будет приходить каждый день.")


async def scheduler_loop():
    log.info("[Бот] Слушатель команд /start запущен...")
    while True:
        now = datetime.now()

        send_time = now.replace(hour=SEND_HOUR, minute=SEND_MINUTE, second=0, microsecond=0)
        if now >= send_time:
            send_time += timedelta(days=1)

        collect_time = send_time - timedelta(minutes=COLLECT_MINUTES_BEFORE)
        if now >= collect_time:
            collect_time += timedelta(days=1)
            send_time = collect_time + timedelta(minutes=COLLECT_MINUTES_BEFORE)

        remaining = (collect_time - now).total_seconds()
        while remaining > 0:
            log.info(f"[Планировщик] Начнём сбор новостей через {remaining / 60:.1f} минут (в {collect_time.strftime('%H:%M')})")
            step = min(60, remaining)
            await asyncio.sleep(step)
            remaining -= step

        log.info("[Планировщик] Время сбора! Начинаю выкачивать каналы...")
        raw_news = await fetch_news_from_channels()
        summary = await get_gemini_summary(raw_news)

        wait_for_send = (send_time - datetime.now()).total_seconds()
        if wait_for_send > 0:
            log.info(f"[Планировщик] Анализ готов. Ждём {wait_for_send:.1f} сек до отправки в {send_time.strftime('%H:%M')}...")
            await asyncio.sleep(wait_for_send)

        await send_summary_to_subscribers(summary)


async def main():
    log.info("Запуск сессий...")
    await client.start()
    await bot_client.start(bot_token=BOT_TOKEN)
    log.info("[Бот] Слушатель команд /start запущен на bot_client...")
    
    await asyncio.gather(
        client.run_until_disconnected(),
        bot_client.run_until_disconnected(),
        scheduler_loop()
    )

if __name__ == "__main__":
    asyncio.run(main())
